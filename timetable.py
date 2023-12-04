import random
import docx
import os

# Function to take inputs for a teacher
def get_teacher_info(subject):
    teacher_info = []
    print("===============================================")
    teacher_info.append(input(f"Enter the teacher for {subject}: "))
    teacher_info.append(input(f"Enter lab room no for {subject}: "))
    return teacher_info

# Number of teachers, days, and periods
# Number of teachers, days, and periods
noofteachers = int(input("Enter the number of teachers: "))
days = 5
periods = 8

# Get the number of subjects
num_subjects = int(input("Enter the number of subjects: "))
subject_names = []

# Get names of each subject
for i in range(num_subjects):
    subject_name = input(f"Enter the name of subject {i + 1}: ")
    subject_names.append(subject_name)
    os.system('cls')

# Manually specify the days for HONOR class
honor_days = input("Enter the days for HONOR class (comma-separated, e.g., 1,3,5): ")
honor_day_list = [int(day) - 1 for day in honor_days.split(',')]

# Manually specify the days for OEC class
oec_days = input("Enter the days for OEC class (comma-separated, e.g., 2,4): ")
oec_day_list = [int(day) - 1 for day in oec_days.split(',')]

# Initialize the timetable
timetable = [[["" for _ in range(periods)] for _ in range(days)] for _ in range(noofteachers)]

# Initialize the timetable
timetable = [[["" for _ in range(periods)] for _ in range(days)] for _ in range(noofteachers)]

# Rest of the code remains the same

subjects = [{"name": subject, "teachers": []} for subject in subject_names]

for subject in subjects:
    for j in range(noofteachers):
        teacher_info = get_teacher_info(subject["name"])
        subject["teachers"].append(teacher_info)
        os.system('cls')

# Assign HONOR and OEC classes
# Assign HONOR and OEC classes
for honor_day in honor_day_list:
    for honor_period in range(periods):
        if honor_period < len(subjects[honor_day]["teachers"]):
            subjects[honor_day]["teachers"][honor_period][0] = "HONOR"

for oec_day in oec_day_list:
    for oec_period in range(periods):
        if oec_period < len(subjects[oec_day]["teachers"]):
            subjects[oec_day]["teachers"][oec_period][0] = "OEC"


# Arrange the remaining timetable
for day in range(days):
    for period in range(periods):
        for subject in subjects:
            if day in honor_day_list and period < len(subject["teachers"][0]) and subject["teachers"][0][period][0] == "HONOR":
                continue
            if day in oec_day_list and period < len(subject["teachers"][0]) and subject["teachers"][0][period][0] == "OEC":
                continue
            subject_assigned = False
            for teacher_info in subject["teachers"]:
                if period < len(teacher_info) and teacher_info[period][0] != "HONOR" and teacher_info[period][0] != "OEC":
                    if timetable[subject["teachers"].index(teacher_info)][day][period] == "":
                        timetable[subject["teachers"].index(teacher_info)][day][period] = subject["name"]
                        subject_assigned = True
                        break
            if subject_assigned:
                break

# Print the timetable to the console
print("Timetable:")
for teacher_idx, teacher_timetable in enumerate(timetable):
    print(f"Teacher {teacher_idx + 1}:")
    for day in range(days):
        for period in range(periods):
            subject = teacher_timetable[day][period]
            if subject:
                print(f"Day {day + 1}, Period {period + 1}: {subject}")

# Save the timetable to a Word document
doc = docx.Document()
doc.add_heading('Timetable', 0)

for teacher_idx, teacher_timetable in enumerate(timetable):
    doc.add_heading(f'Teacher {teacher_idx + 1}', level=1)
    table = doc.add_table(rows=days, cols=periods)
    for day in range(days):
        for period in range(periods):
            cell = table.cell(day, period)
            subject = teacher_timetable[day][period]
            cell.text = subject

doc.save('timetable.docx')
print("Timetable saved to timetable.docx")

